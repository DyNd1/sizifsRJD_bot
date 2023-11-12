import re
from docx import Document
import pandas as pd

def is_start_new_item(s):
    # Если в начале строки видим число, точку и пробел, то это начало пункта
    pattern = re.compile(r'^\d+\. ')
    if pattern.findall(s.strip()):
        return True
    else:
        return False


def is_bad_paragraph(s):
    pattern1 = re.compile(r'^<\d+> ')
    pattern2 = re.compile(r'^[IXVLM]')

    # Если в начале строки видим <, число, > и пробел то это говно подпункт
    if pattern1.findall(s.strip()):
        return True
    # а это залупный разделитель, тоже нахуй
    elif '--------------------------------' in s:
        return True
    # Если в начале строки видим римские цифры, то это заголовок, убираем нахуй
    elif pattern2.findall(s.strip()):
        return True
    else:
        return False


# Итоговый массив с пунктами
items = []

# Я изменил документ, убрав говно снизу
document = Document("ржд.docx")

# Временный массив, для хранения абзацев пункта (в одном пункте могут быть несколько абзацев)
temp = []

# Флаг, показывает собираем мы сейчас пункт или нет
flag = False

# Перебор абзацев
for p in document.paragraphs:
    p = p.text

    if is_bad_paragraph(p):
        # Прекращаем сбор пункта если встречаем плохие абзаца типа
        # -----------,
        # <1> Абзацы третий, четвертый пункта 5 Регламента действий локомот...
        # II. Сигналы на железнодорожном транспорте
        flag = False

    elif p.strip().startswith("Приложение N "):
        # Прекращаем сбор пункта если встречаем начало приложения
        # Внимание! ПУНКТЫ из приложений собираются, так как посчитал что они нужны, но не собираем заголовки
        # и названия разделов
        flag = False

    elif is_start_new_item(p):
        # Если видим что начинается новый пункт, то прошлый добавляем в аутпут массив
        items.append(' '.join(temp))
        # Чистим временное хранилище в котором хранятся абзацы пунктов
        temp = []
        # Добавляем начало нового пункта во временное хранилище
        temp.append(p)
        # Ставим флаг, что начался новый пункт
        flag = True

    elif flag:
        # Если флаг нового пункта Тру, то просто добавляем текущие абзацы в наш пункт
        temp.append(p)

# Добавляем последний абзац, чтобы заебок было
items.append(' '.join(temp))

# Убираем пустые пункты и убираем говно из начала и конца каждого пункта
items = [s.strip() for s in items if s]

print(items[100])

print(f"Собранное количество пунктов: {len(items)}")

# Здесь проверяю количество пунктов (должно совпадать по идеи с предыдущим, у меня совпало)
pattern = re.compile(r"\n\d+\. ", )
text = ' '.join([p.text for p in document.paragraphs])
matches = pattern.findall(text)
print(f"Проверочное количество пунктов: {len(matches)}")
generator_items = (x for x in items)
df = pd.DataFrame({"context":  generator_items})

df.to_csv("databasetest.csv")

